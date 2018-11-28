'''
Created on Nov 27, 2018

@author: ashish.maikhuri
'''

import Excle 
import os
from docx  import Document
os.chdir('D:\\test')
doc=Document('test.docx')

def table1(tbl):
    for i,row in enumerate(tbl.rows,1):
        if(i==2):
            temp=Excle.awatinguserinfotable()
        elif(i==3):
            temp=Excle.assignedtable()
        elif(i==4):
            temp=Excle.wiptable()
        elif(i==5):
            temp=Excle.aw3rdpartytable()
        elif(i==6):
            temp=Excle.awchangetable()
        for j,cell in enumerate(row.cells,1) :
            if not (i==1 or j==1):
                cell.vlaue=temp[j-1]
                
def table2(tbl):
    temp=Excle.reqtable()
    for i,rows in enumerate(tbl.row,1):
        for j,cell in enumerate(rows.cells,1):
            if(i==1 or j==1):
                continue
            else:
                cell.value=temp[i-2]
    


    
for s,i in enumerate(doc.paragraphs,1):
    Excle.datafill()
    if(s==5):
        temp=Excle.totalincident()
        i.text=temp
    elif(s==6):
        temp=Excle.AUFincident()
        i.text=temp
    elif(s==7):
        temp=Excle.Assignedincident()
        i.text=temp
    elif(s==8):
        temp=Excle.wipincident()
        i.text=temp
    elif(s==9):
        temp=Excle.A3partyincident()
        i.text=temp
    elif(s==10):
        temp=Excle.Achangeincident()
        i.text=temp
    elif(s==13):
        temp=Excle.allreq()
        i.text=temp
for s,tb in enumerate(doc.tables,1):
    if(s==1):
        table1(tb)
    elif(s==2):
        table2(tb)
doc.save('Queuestatus.docx')
